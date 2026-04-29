import os
import tempfile
from docx import Document

def fill_docx_template(template_path, data_dict):
    """
    Reads a .docx template, replaces {{key}} with values from data_dict,
    and returns the physical path to the temporarily saved output file.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Missing template: {template_path}")
        
    doc = Document(template_path)

    # Replace in regular paragraphs
    for para in doc.paragraphs:
        for key, value in data_dict.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value))

    # Replace in tables (which contain their own paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in data_dict.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, str(value))

    # Save to a temporary file and return the path
    temp_dir = tempfile.gettempdir()
    # E.g. "C:/Temp/filled_agreement.docx"
    output_path = os.path.join(temp_dir, f"filled_agreement_{os.urandom(4).hex()}.docx")
    doc.save(output_path)
    
    return output_path
